from stemmer import TurkishStemmer
from normalizer import TurkishNormalizer
from stopword import TurkishStopwordRemover
from ml_based_tokenizer import MlBasedTokenizer
from ml_based_splitter import MlBasedSentenceSplitter
from rule_based_tokenizer import RuleBasedTokenizer
from rule_based_splitter import RuleBasedSentenceSplitter

import argparse
import os

# Importing docx library for Word documents
try:
    from docx import Document
    permit_word_file = True
    
except ImportError as e:
    permit_word_file = False
       
    
parser = argparse.ArgumentParser(description="Preprocessing documents written in Turkish language...")

parser.add_argument("-files", nargs="+", help="path for the files to be processed")
parser.add_argument('-grammar_path', default="./data/grammar.json", type=str, help="path of the grammar file for the stemmer")
parser.add_argument('-lexicon_path', default="./data/lexicon.txt", type=str, help="path of the lexicon file")
parser.add_argument('-corpus_path', default="./data/corpus.txt", type=str, help="path of the corpus file for language model in the normalizer")
parser.add_argument('-train_corpus_path', default="./data/train.conllu", type=str, help="path of the corpus file for training ml based tokenizer and sentence splitter")
parser.add_argument('-stopword_path', default="./data/stopword_lexicon.txt", type=str, help="path of the static stopword lexicon file")

parser.add_argument('-ngram', default=3, type=int,   help="ngram for the language model in the normalizer")
parser.add_argument('-do_train', default=False, help="train the ml based tokenizer and sentence splitter")

parser.add_argument('-stem', help="apply stemming in the given files", action='store_true')
parser.add_argument('-normalize', help="apply normalization in the given files", action='store_true')
parser.add_argument('-stopword', help="remove stopwords in the given files", action='store_true')
parser.add_argument('-run_stopword_analysis', help="run stopword analysis on the given corpus file", action='store_true')
parser.add_argument('-ml_tokenize', help="tokenize the given files using ml based tokenizer", action='store_true')
parser.add_argument('-rule_tokenize', help="tokenize the given files using rule based tokenizer", action='store_true')
parser.add_argument('-ml_split', help="split the given files using ml based sentence splitter", action='store_true')
parser.add_argument('-rule_split', help="split the given files using rule based sentence splitter", action='store_true')
parser.add_argument('-interact', help="starts interactive application", action='store_true')
parser.add_argument('-save', help="saving the files according to the given format", action='store_true')



args = parser.parse_args()


def save_file(pro_str: str, path: str, fname: str) -> None:
    if "docx" in fname and permit_word_file:
        doc = Document()
        doc.add_paragraph().add_run(pro_str)
        doc.save(f"{path}/processed_{fname}")
        
    else:
        with open(f"{path}/processed_{fname}", "w", encoding="utf8") as file:
            file.write(pro_str)

            
def read_file(path: str, fname: str) -> [str]:
    if ".docx" in fname and permit_word_file:
        doc = Document(f"{path}/{fname}")
        pre_str = [sent.text for sent in doc.paragraphs]
        
    elif ".txt" in fname:
        with open(f"{path}/{fname}", "r", encoding="utf8") as file:
            pre_str = file.read().split("\n")      
            
    else: pre_str = -1
           
    return pre_str


def interact(normalizer: TurkishNormalizer, stemmer: TurkishStemmer, stopwordRemover: TurkishStopwordRemover, mlBasedTokenizer: MlBasedTokenizer, ruleBasedTokenizer:RuleBasedTokenizer, mlBasedSplitter: MlBasedSentenceSplitter, ruleBasedSplitter: RuleBasedSentenceSplitter) -> None:
    menu = "\t[mt]lBasedTokenize\n\t[rt]uleBasedTokenize\n\t[ms]lBasedSplit\n\t[rs]uleBasedSplit\n\t[n]ormalize\n\t[s]tem\n\tstop[w]ord\n\t[a]ll\n\t[e]xit\n>> Selection: "
    opt = input(menu)
    
    while(opt != "e"):

        if opt == "mt":
            sentence = input(">> Sentence: ")
            print(f":: Tokenized sent: {mlBasedTokenizer.tokenize(sentence)}")
        
        elif opt == "rt":
            sentence = input(">> Sentence: ")
            print(f":: Tokenized sent: {ruleBasedTokenizer.tokenize(sentence)}")
        
        elif opt == "ms":
            sentence = input(">> Sentence: ")
            print(f":: Splitted sent: {mlBasedSplitter.split(sentence)}")
        
        elif opt == "rs":
            sentence = input(">> Sentence: ")
            print(f":: Splitted sent: {ruleBasedSplitter.split(sentence)}")

        elif opt == "n":
            sentence = input(">> Sentence: ")
            print(f":: Normalized sent: {normalizer.normalize_sentence(sentence)}")
            
        elif opt == "s":
            sentence = input(">> Sentence: ")
            print(f":: Stemmed sent: {stemmer.stem_sentence(sentence)}")
            
        elif opt == "w":
            sentence = input(">> Sentence: ")
            print(f":: Stopword removed sent: {stopwordRemover.remove_stopwords(sentence)}")
            
        elif opt == "a":
            sentence = input(">> Sentence: ")
            ml_tokenized = mlBasedTokenizer.tokenize(sentence)
            rule_tokenized = ruleBasedTokenizer.tokenize(sentence)
            ml_splitted = mlBasedSplitter.split(sentence)
            rule_splitted = ruleBasedSplitter.split(sentence)
            sentence = normalizer.normalize_sentence(sentence)
            sentence = stemmer.stem_sentence(sentence)
            sentence = stopwordRemover.remove_stopwords(sentence)
            
            print(f":: Machine learning based tokenized sent: {ml_tokenized}")
            print(f":: Rule based tokenized sent: {rule_tokenized}")
            print(f":: Machine learning based splitted sent: {ml_splitted}")
            print(f":: Rule based splitted sent: {rule_splitted}")
            print(f":: Preprocessed sent: {sentence}")
        
        else: 
            print("xx Invalid input...")
        
        opt = input(menu)
            

if __name__ == "__main__":
    mlTokenizer = MlBasedTokenizer(args.train_corpus_path, args.do_train)
    ruleTokenizer = RuleBasedTokenizer()
    mlSplitter = MlBasedSentenceSplitter(args.train_corpus_path, args.do_train)
    ruleSplitter = RuleBasedSentenceSplitter()
    normalizer = TurkishNormalizer(args.corpus_path, args.lexicon_path, ngram = args.ngram)
    stemmer    = TurkishStemmer(args.lexicon_path, args.grammar_path)
    stopwordRemover = TurkishStopwordRemover(args.stopword_path, args.corpus_path, use_dynamic=True)
    
    if args.interact:
        interact(normalizer, stemmer, stopwordRemover, mlTokenizer, ruleTokenizer, mlSplitter, ruleSplitter)
        
    elif args.run_stopword_analysis:
        stopwordRemover.print_analysis()
        stopwordRemover.plot_analysis()
        
    else:
        for fname in args.files:
            # extracting the path of the given file
            path = fname.split("/")
            if os.path.isdir(path[0]):
                path, fname = path[0], path[1]
            else:
                path, fname = "./", fname

            # reading the file
            pre_str = read_file(path, fname)

            if pre_str == -1:
                print(f">> Given the file named {fname} is in unsupported format...")
                continue

            # applying related functionalities
            if args.stem and not args.normalize : 
                pro_str = [stemmer.stem_sentence(sent) for sent in pre_str]

            elif not args.stem and args.normalize:
                pro_str = [normalizer.normalize_sentence(sent) for sent in pre_str]
            
            elif args.ml_tokenize:
                pro_str = [mlTokenizer.tokenize(sent) for sent in pre_str]
                pro_str = [" ".join(sent) for sent in pro_str]
            
            elif args.rule_tokenize:
                pro_str = [ruleTokenizer.tokenize(sent) for sent in pre_str]
                pro_str = [" ".join(sent) for sent in pro_str]
            
            elif args.ml_split:
                pro_str = [mlSplitter.split(sent) for sent in pre_str]
                pro_str = ["\n".join(sent) for sent in pro_str]
            
            elif args.rule_split:
                pro_str = [ruleSplitter.split(sent) for sent in pre_str]
                pro_str = ["\n".join(sent) for sent in pro_str]

            else:
                pro_str = [normalizer.normalize_sentence(sent) for sent in pre_str]
                pro_str = [stemmer.stem_sentence(sent) for sent in pro_str]
                pro_str = [stopwordRemover.remove_stopwords(sent) for sent in pro_str]
            
            # saving the file
            if args.save is not None:
                save_file("\n".join(pro_str), path, fname)
